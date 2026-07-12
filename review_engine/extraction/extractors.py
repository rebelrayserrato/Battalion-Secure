from __future__ import annotations

import logging
import os
import re
import tempfile
import zipfile
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document

from review_engine.config.settings import (
    CHUNK_OVERLAP,
    CHUNK_SIZE,
    IMAGE_EXTENSIONS,
    OCR_DPI,
    OCR_MIN_NATIVE_CHARS,
    SUPPORTED_EXTENSIONS,
    ZIP_MAX_FILES,
    ZIP_MAX_RATIO,
    ZIP_MAX_TOTAL_BYTES,
)
from review_engine.extraction import ocr
from review_engine.extraction.models import SourceChunk, source_reference

logger = logging.getLogger(__name__)


def _split_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    clean = re.sub(r"[ \t]+", " ", text).strip()
    if not clean:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(clean):
        end = min(len(clean), start + size)
        if end < len(clean):
            boundary = max(clean.rfind("\n", start, end), clean.rfind(". ", start, end))
            if boundary > start + size // 2:
                end = boundary + 1
        chunks.append(clean[start:end].strip())
        if end >= len(clean):
            break
        start = max(start + 1, end - overlap)
    return chunks


def _make_chunks(
    matter_id: str,
    path: Path,
    texts: Iterable[tuple[str, int | None, int | None, str | None]],
) -> list[SourceChunk]:
    result: list[SourceChunk] = []
    ordinal = 0
    for text, page, row, section in texts:
        for part in _split_text(text):
            result.append(
                SourceChunk(
                    matter_id=matter_id,
                    document_name=path.name,
                    file_type=path.suffix.lower().lstrip("."),
                    page=page,
                    row=row,
                    section=section,
                    text=part,
                    source_ref=source_reference(
                        matter_id,
                        path.name,
                        page=page,
                        row=row,
                        section=section,
                        ordinal=ordinal,
                    ),
                )
            )
            ordinal += 1
    return result


def _extract_pdf(path: Path) -> list[tuple[str, int | None, int | None, str | None]]:
    import fitz

    output = []
    with fitz.open(path) as pdf:
        for number, page in enumerate(pdf, start=1):
            text = page.get_text("text")
            tables = []
            try:
                for table in page.find_tables().tables:
                    tables.append("\n".join(" | ".join(str(v or "") for v in row) for row in table.extract()))
            except (AttributeError, TypeError, ValueError):
                pass
            combined = text + ("\nTABLE:\n" + "\n".join(tables) if tables else "")
            section = None
            # OCR fallback: only for pages whose native text layer is empty or
            # near-empty (scanned / image-only pages). Native pages are untouched.
            if len(text.strip()) < OCR_MIN_NATIVE_CHARS and ocr.ocr_available():
                ocr_text = _ocr_pdf_page(page)
                if ocr_text:
                    combined = (combined.strip() + "\n" + ocr_text).strip() if combined.strip() else ocr_text
                    section = "OCR text"
            output.append((combined, number, None, section))
    return output


def _ocr_pdf_page(page) -> str:
    try:
        pixmap = page.get_pixmap(dpi=OCR_DPI)
        return ocr.ocr_png_bytes(pixmap.tobytes("png"))
    except Exception as exc:  # pragma: no cover - render failure is non-fatal
        logger.warning("Could not render PDF page for OCR (%s)", exc)
        return ""


def _extract_image(path: Path) -> list[tuple[str, int | None, int | None, str | None]]:
    text = ocr.ocr_image_file(path)
    if not text:
        return []
    return [(text, None, None, "OCR text")]


def _extract_docx(path: Path) -> list[tuple[str, int | None, int | None, str | None]]:
    doc = Document(path)
    output = []
    current_section = "Document body"
    buffer: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            if buffer:
                output.append(("\n".join(buffer), None, None, current_section))
                buffer = []
            current_section = text
        else:
            buffer.append(text)
    if buffer:
        output.append(("\n".join(buffer), None, None, current_section))
    for index, table in enumerate(doc.tables, start=1):
        text = "\n".join(" | ".join(cell.text for cell in row.cells) for row in table.rows)
        output.append((text, None, None, f"Table {index}"))
    return output


def _extract_spreadsheet(path: Path) -> list[tuple[str, int | None, int | None, str | None]]:
    sheets = (
        pd.read_excel(path, sheet_name=None, dtype=str)
        if path.suffix.lower() == ".xlsx"
        else {"CSV": pd.read_csv(path, dtype=str, keep_default_na=False)}
    )
    output = []
    for sheet_name, frame in sheets.items():
        frame = frame.fillna("")
        for position, (_, record) in enumerate(frame.iterrows(), start=2):
            text = " | ".join(f"{column}: {record[column]}" for column in frame.columns)
            output.append((text, None, position, str(sheet_name)))
    return output


def _safe_zip_member_name(member: str) -> str | None:
    """Return a flattened, traversal-safe filename for a zip member, or None
    if the member path is unsafe (absolute or escapes the archive root)."""
    normalized = os.path.normpath(member.replace("\\", "/"))
    parts = Path(normalized).parts
    if os.path.isabs(normalized) or (parts and (parts[0] == ".." or parts[0].endswith(":"))):
        return None
    if any(part == ".." for part in parts):
        return None
    flattened = "_".join(parts)
    return re.sub(r"[^A-Za-z0-9._-]", "_", flattened).strip("_") or None


def _extract_zip(path: Path, matter_id: str) -> list[SourceChunk]:
    """Unpack a ZIP safely and route each contained file through the existing
    extractor dispatch. Guards against path traversal and zip bombs (per-file
    ratio, total uncompressed size, and member count)."""
    chunks: list[SourceChunk] = []
    total_bytes = 0
    processed = 0
    with zipfile.ZipFile(path) as archive:
        for info in archive.infolist():
            if info.is_dir():
                continue
            member = info.filename
            extension = Path(member).suffix.lower()
            # Never recurse into nested archives — a cheap zip-bomb vector.
            if extension == ".zip":
                logger.info("ZIP %s: skipping nested archive %s", path.name, member)
                continue
            if extension not in SUPPORTED_EXTENSIONS:
                continue
            safe_name = _safe_zip_member_name(member)
            if safe_name is None:
                logger.warning("ZIP %s: skipping unsafe member path %s", path.name, member)
                continue
            if info.compress_size > 0 and info.file_size / info.compress_size > ZIP_MAX_RATIO:
                raise ValueError(
                    f"ZIP {path.name}: member {member} exceeds compression ratio guard"
                )
            processed += 1
            if processed > ZIP_MAX_FILES:
                raise ValueError(f"ZIP {path.name}: too many files (>{ZIP_MAX_FILES})")
            total_bytes += info.file_size
            if total_bytes > ZIP_MAX_TOTAL_BYTES:
                raise ValueError(f"ZIP {path.name}: uncompressed size exceeds guard")
            with tempfile.TemporaryDirectory() as tmp:
                # Provenance-preserving name: "<zip stem>__<flattened member>".
                target = Path(tmp) / f"{path.stem}__{safe_name}"
                with archive.open(info) as src, open(target, "wb") as dst:
                    remaining = info.file_size
                    while True:
                        block = src.read(65536)
                        if not block:
                            break
                        remaining -= len(block)
                        if remaining < 0:
                            raise ValueError(
                                f"ZIP {path.name}: member {member} larger than declared size"
                            )
                        dst.write(block)
                chunks.extend(extract_document(target, matter_id))
    return chunks


def extract_document(path: str | Path, matter_id: str) -> list[SourceChunk]:
    path = Path(path)
    extension = path.suffix.lower()
    if extension == ".zip":
        return _extract_zip(path, matter_id)
    if extension == ".pdf":
        texts = _extract_pdf(path)
    elif extension == ".docx":
        texts = _extract_docx(path)
    elif extension in {".csv", ".xlsx"}:
        texts = _extract_spreadsheet(path)
    elif extension == ".txt":
        texts = [(path.read_text(encoding="utf-8", errors="replace"), None, None, "Document body")]
    elif extension in IMAGE_EXTENSIONS:
        texts = _extract_image(path)
    else:
        raise ValueError(f"Unsupported file type: {extension}")
    return _make_chunks(matter_id, path, texts)
