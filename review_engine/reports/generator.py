from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)
from xml.sax.saxutils import escape

from review_engine.reports.decisions import (
    APPROVED,
    NEEDS_CHANGES,
    REJECTED,
    UNDECIDED,
    finding_decision,
    load_decisions,
    rollup,
)

# RAYSERR identity palette (matches admin.css / RAYAAAA-227 review-engine theme).
NAVY = "#1b2f5b"
GOLD = "#c8922a"
TEXT = "#1a1f2e"
WORDMARK = "RAYSERR SOLUTIONS"
TAGLINE = "Evidence Review Engine"

_NAVY_RGB = RGBColor(0x1B, 0x2F, 0x5B)
_GOLD_RGB = RGBColor(0xC8, 0x92, 0x2A)
_TEXT_RGB = RGBColor(0x1A, 0x1F, 0x2E)

_STATUS_LABELS = {
    APPROVED: "Approved",
    REJECTED: "Rejected",
    NEEDS_CHANGES: "Needs changes",
    UNDECIDED: "Undecided",
}

SECTIONS = [
    "Executive Summary",
    "Documents Reviewed",
    "Timeline",
    "Key People and Entities",
    "Missing Documents",
    "HR / Legal Risk Flags",
    "Fraud Red Flags",
    "Contradictions",
    "Unsupported or Low-Confidence Findings",
    "Human Review Required",
    "Reviewer Decisions",
    "Recommended Next Steps",
    "Source Appendix",
    "Audit Log Summary",
]


def _generation_timestamp(generated_at: datetime | str | None) -> str:
    """Deterministic-friendly timestamp: honour a caller-supplied value, else now (UTC)."""
    if generated_at is None:
        generated_at = datetime.now(timezone.utc)
    if isinstance(generated_at, datetime):
        return generated_at.strftime("%Y-%m-%d %H:%M UTC")
    return str(generated_at)


def _rollup_sentence(counts: dict) -> str:
    return (
        f"Reviewer decisions: {counts[APPROVED]} approved, {counts[REJECTED]} rejected, "
        f"{counts[NEEDS_CHANGES]} needs-changes, {counts[UNDECIDED]} undecided."
    )


def _report_data(
    db,
    matter_id: str,
    executive_summary: str | None = None,
    decisions=None,
    generated_at: datetime | str | None = None,
) -> dict:
    matter = db.get_matter(matter_id)
    if not matter:
        raise ValueError(f"Unknown matter: {matter_id}")
    documents = db.list_documents(matter_id)
    findings = db.get_findings(matter_id)
    chunks = db.get_chunks(matter_id)
    from review_engine.evidence.timeline import build_timeline

    timeline = build_timeline(chunks)
    entities = db.get_entities(matter_id)
    audits = db.get_audit_log(matter_id)

    decisions = load_decisions(decisions, matter_id)
    counts = rollup(findings, decisions) if decisions else {}

    summary = executive_summary or (
        f"This evidence review identified {len(findings)} source-supported review flags "
        f"across {len(documents)} document(s). Findings are screening indicators only, "
        "not legal conclusions or determinations that fraud occurred."
    )
    if decisions:
        summary = f"{summary} {_rollup_sentence(counts)}"

    grouped = {
        "Missing Documents": [f for f in findings if f["category"] == "Missing Document"],
        "HR / Legal Risk Flags": [f for f in findings if f["category"] == "HR Legal Risk"],
        "Fraud Red Flags": [f for f in findings if f["category"] == "Fraud Red Flag"],
        "Contradictions": [f for f in findings if f["category"] == "Contradiction"],
        "Unsupported or Low-Confidence Findings": [
            f for f in findings if f["category"] == "Unsupported Finding" or f["confidence"] == "Low"
        ],
        "Human Review Required": [f for f in findings if f["human_review_required"]],
    }
    return {
        "matter": matter, "documents": documents, "findings": findings, "chunks": chunks,
        "timeline": timeline, "entities": entities, "audits": audits, "summary": summary,
        "grouped": grouped, "decisions": decisions, "decision_counts": counts,
        "generated_at": _generation_timestamp(generated_at),
    }


def _finding_text(finding: dict) -> str:
    citations = "; ".join(source["citation"] for source in finding["supporting_sources"])
    return (
        f"{finding['title']} [{finding['confidence']} confidence] — {finding['explanation']} "
        f"Confidence basis: {finding['confidence_reason']} Sources: {citations}"
    )


def _decision_text(finding: dict, src_id: str, decision: dict) -> str:
    label = _STATUS_LABELS.get(decision["status"], decision["status"])
    parts = [f"{finding['title']} — {label} [{src_id}]"]
    if decision.get("note"):
        parts.append(f"Note: {decision['note']}")
    attribution = []
    if decision.get("reviewer"):
        attribution.append(f"reviewer {decision['reviewer']}")
    if decision.get("decided_at"):
        attribution.append(str(decision["decided_at"]))
    if attribution:
        parts.append(f"({', '.join(attribution)})")
    return " ".join(parts)


def _section_items(data: dict, section: str) -> list[str]:
    if section == "Executive Summary":
        return [data["summary"]]
    if section == "Documents Reviewed":
        return [
            f"{doc['name']} ({doc['file_type']}); processed: {doc['processed_at'] or 'not processed'}"
            for doc in data["documents"]
        ] or ["No documents."]
    if section == "Timeline":
        return [
            f"{item['date']} — {item['event']} Source: {item['citation']}"
            for item in data["timeline"]
        ] or ["No dated events identified."]
    if section == "Key People and Entities":
        return [
            f"{entity['entity_type']}: {entity['value']} ({entity['source_ref']})"
            for entity in data["entities"]
            if entity["entity_type"] != "event"
        ] or ["No entities identified."]
    if section == "Reviewer Decisions":
        decisions = data["decisions"]
        if not decisions:
            return []
        items = []
        for finding in data["findings"]:
            src_id, decision = finding_decision(finding, decisions)
            if decision is None:
                continue
            items.append(_decision_text(finding, src_id, decision))
        items.append(_rollup_sentence(data["decision_counts"]))
        return items
    if section in data["grouped"]:
        return [_finding_text(f) for f in data["grouped"][section]] or ["No source-supported findings."]
    if section == "Recommended Next Steps":
        return [
            "Have a qualified human reviewer validate each cited source and resolve contradictions.",
            "Confirm document-set completeness and the applicable jurisdiction.",
            "Do not treat anomaly scores or red flags as proof of wrongdoing.",
        ]
    if section == "Source Appendix":
        return [f"{chunk.source_ref}: {chunk.citation} — {chunk.text[:300]}" for chunk in data["chunks"]]
    if section == "Audit Log Summary":
        return [
            f"{entry['timestamp']} — {entry['event_type']}: {entry['details'] or ''}"
            for entry in data["audits"]
        ] or ["No audit entries."]
    return []


def _active_sections(data: dict) -> list[str]:
    """Skip the Reviewer Decisions section entirely when no decisions are present (AC3)."""
    if data["decisions"]:
        return SECTIONS
    return [s for s in SECTIONS if s != "Reviewer Decisions"]


def generate_docx_report(
    db,
    matter_id: str,
    output: str | Path | None = None,
    executive_summary: str | None = None,
    decisions=None,
    generated_at: datetime | str | None = None,
) -> bytes:
    data = _report_data(db, matter_id, executive_summary, decisions, generated_at)
    document = Document()

    wordmark = document.add_paragraph()
    run = wordmark.add_run(WORDMARK)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = _NAVY_RGB
    tag = wordmark.add_run(f"   {TAGLINE}")
    tag.font.size = Pt(11)
    tag.font.color.rgb = _GOLD_RGB

    heading = document.add_heading("Evidence Review Report", 0)
    for hrun in heading.runs:
        hrun.font.color.rgb = _NAVY_RGB

    meta = document.add_paragraph()
    meta.add_run(f"Task: {data['matter']['name']} ({matter_id})\n").bold = True
    meta.add_run(f"Generated: {data['generated_at']}\n").font.color.rgb = _TEXT_RGB
    meta.add_run(
        f"Jurisdiction: {data['matter']['jurisdiction'] or 'Required / not specified'}"
    ).font.color.rgb = _TEXT_RGB

    document.add_paragraph(
        "Human-review document. This report does not provide legal advice and does not conclude that fraud occurred."
    )
    for section in _active_sections(data):
        section_heading = document.add_heading(section, level=1)
        for hrun in section_heading.runs:
            hrun.font.color.rgb = _NAVY_RGB
        for item in _section_items(data, section):
            document.add_paragraph(item, style="List Bullet")
    buffer = BytesIO()
    document.save(buffer)
    content = buffer.getvalue()
    if output:
        Path(output).write_bytes(content)
    return content


def generate_pdf_report(
    db,
    matter_id: str,
    output: str | Path | None = None,
    executive_summary: str | None = None,
    decisions=None,
    generated_at: datetime | str | None = None,
) -> bytes:
    data = _report_data(db, matter_id, executive_summary, decisions, generated_at)
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=LETTER, rightMargin=0.65 * inch, leftMargin=0.65 * inch,
        topMargin=0.65 * inch, bottomMargin=0.65 * inch,
        # invariant=1 pins reportlab's creation date / document id so a given input
        # yields byte-identical output (AC4 determinism, headless/offline).
        invariant=1,
    )
    styles = getSampleStyleSheet()
    navy = colors.HexColor(NAVY)
    gold = colors.HexColor(GOLD)
    body_color = colors.HexColor(TEXT)

    wordmark_style = ParagraphStyle(
        "Wordmark", parent=styles["Normal"], fontName="Helvetica-Bold",
        fontSize=18, textColor=navy, spaceAfter=2,
    )
    tagline_style = ParagraphStyle(
        "Tagline", parent=styles["Normal"], fontName="Helvetica",
        fontSize=10, textColor=gold, spaceAfter=6,
    )
    title_style = ParagraphStyle(
        "ReportTitle", parent=styles["Title"], textColor=navy, alignment=TA_CENTER,
    )
    heading_style = ParagraphStyle(
        "SectionHeading", parent=styles["Heading1"], textColor=navy,
    )
    body_style = ParagraphStyle(
        "ReportBody", parent=styles["BodyText"], textColor=body_color,
    )
    meta_style = ParagraphStyle(
        "Meta", parent=styles["Normal"], textColor=body_color,
    )

    story = [
        Paragraph(WORDMARK, wordmark_style),
        Paragraph(TAGLINE, tagline_style),
        HRFlowable(width="100%", thickness=1.5, color=gold, spaceBefore=2, spaceAfter=10),
        Paragraph("Evidence Review Report", title_style),
        Spacer(1, 8),
        Paragraph(escape(f"Task: {data['matter']['name']} ({matter_id})"), meta_style),
        Paragraph(escape(f"Generated: {data['generated_at']}"), meta_style),
        Paragraph(
            escape(f"Jurisdiction: {data['matter']['jurisdiction'] or 'Required / not specified'}"),
            meta_style,
        ),
        Paragraph(
            "Human-review document. This report does not provide legal advice and does not conclude that fraud occurred.",
            styles["Italic"],
        ),
        Spacer(1, 12),
    ]
    for section in _active_sections(data):
        story.append(Paragraph(escape(section), heading_style))
        for item in _section_items(data, section):
            story.append(Paragraph("• " + escape(item), body_style))
            story.append(Spacer(1, 5))
    doc.build(story)
    content = buffer.getvalue()
    if output:
        Path(output).write_bytes(content)
    return content
